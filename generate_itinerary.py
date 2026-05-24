#!/usr/bin/env python3
"""Generate an editable itinerary DOCX from an order spreadsheet."""

from __future__ import annotations

import argparse
import json
import re
import shutil
import sys
import unicodedata
import warnings
from datetime import date, datetime
from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from openpyxl import load_workbook


warnings.filterwarnings(
    "ignore",
    message="Workbook contains no default style.*",
    category=UserWarning,
    module="openpyxl.styles.stylesheet",
)

BASE_DIR = Path(__file__).resolve().parent
INPUT_DIR = BASE_DIR / "input"
TEMPLATE_DIR = BASE_DIR / "templates"
OUTPUT_DIR = BASE_DIR / "output"
CONFIG_DIR = BASE_DIR / "config"
DEFAULT_TEMPLATE = TEMPLATE_DIR / "行程单通用模版.docx"
DEFAULT_CONFIG = CONFIG_DIR / "trip_config.json"
ORDER_PATTERN = "订单列表_*.xlsx"
SHEET_NAME = "订单数据"
REQUIRED_COLUMNS = ("旅客姓名", "旅客证件类型", "旅客证件号码", "旅客手机号码")
OPTIONAL_ORDER_COLUMNS = ("路线", "出行时间")
DEFAULT_TRIP_FIELDS = {
    "agency": "广州行走贰拾国际旅行社有限公司",
    "source": "广东",
    "start_date": "",
    "end_date": "",
    "operator": "焦一航",
    "operator_phone": "15915214565",
}


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(description="生成行程单 DOCX")
    parser.add_argument("--excel", type=Path, help="订单 Excel 路径，默认读取当前目录最新的订单列表_*.xlsx")
    parser.add_argument("--template", type=Path, default=DEFAULT_TEMPLATE, help="DOCX 模板路径")
    parser.add_argument("--config", type=Path, default=DEFAULT_CONFIG, help="JSON 配置路径")
    parser.add_argument("--output", type=Path, help="输出 DOCX 路径，默认使用配置里的 output_name")
    parser.add_argument("--no-prompt", action="store_true", help="不在终端询问，直接使用配置文件")
    return parser.parse_args()


def fail(message: str) -> None:
    raise SystemExit(f"错误：{message}")


def resolve_path(path: Path) -> Path:
    return path if path.is_absolute() else BASE_DIR / path


def latest_order_file() -> Path:
    search_dirs = (BASE_DIR, INPUT_DIR)
    files = sorted(
        [path for directory in search_dirs for path in directory.glob(ORDER_PATTERN)],
        key=lambda p: p.stat().st_mtime,
        reverse=True,
    )
    if not files:
        fail(f"没有找到订单表格，请把 {ORDER_PATTERN} 放到项目根目录或 input 文件夹")
    return files[0]


def load_config(path: Path) -> dict[str, Any]:
    if not path.exists():
        fail(f"找不到配置文件：{path}")
    try:
        with path.open("r", encoding="utf-8") as f:
            config = json.load(f)
    except json.JSONDecodeError as exc:
        fail(f"配置文件不是合法 JSON：{path}（第 {exc.lineno} 行第 {exc.colno} 列）")

    config = {**DEFAULT_TRIP_FIELDS, **config}
    required = (
        "route",
        "leaders",
        "leader_phone",
    )
    missing = [key for key in required if not str(config.get(key, "")).strip()]
    if missing:
        fail("配置文件缺少必填字段：" + "、".join(missing))
    return config


def prompt_trip_fields(config: dict[str, Any]) -> dict[str, Any]:
    while True:
        print("\n请填写行程信息。直接按回车会使用默认值。粘贴内容如果自带换行，会自动提交。\n")
        updated = dict(config)
        updated["route"] = prompt_value("旅游路线", updated.get("route", ""), show_default=False)
        print("已记录旅游路线。\n")
        updated["leaders"] = prompt_value(
            "领队",
            updated.get("leaders", ""),
            show_default=False,
            hint="多领队用、分隔",
        )
        updated["leader_phone"] = prompt_value(
            "领队电话",
            updated.get("leader_phone", ""),
            show_default=False,
            hint="填写手机号",
        )

        if confirm_trip_fields(updated):
            print()
            return updated

        print("\n已取消本次填写，请重新输入。")


def prompt_value(label: str, default: Any, show_default: bool = True, hint: str = "") -> str:
    default_text = str(default or "").strip()
    prompt = f"{label}"
    if default_text and show_default:
        prompt += f" [{default_text}]"
    elif hint:
        prompt += f"（{hint}，回车使用默认值）"
    elif default_text:
        prompt += "（回车使用默认值）"
    prompt += "，输入完成后按回车"

    while True:
        print(prompt)
        try:
            value = input("> ").strip()
        except EOFError:
            fail("已退出，未生成行程单")
        print()
        if value:
            return value
        if default_text:
            return default_text
        print(f"{label}不能为空，请重新输入。")


def confirm_trip_fields(config: dict[str, Any]) -> bool:
    print("\n请确认信息：")
    print_wrapped_field("旅游路线", config["route"])
    print_wrapped_field("领队", config["leaders"])
    print_wrapped_field("领队电话", config["leader_phone"])
    if config.get("start_date") and config.get("end_date"):
        print_wrapped_field("出行时间", f"{format_chinese_date(config['start_date'])} 至 {format_chinese_date(config['end_date'])}")
    if config.get("route_name"):
        print_wrapped_field("表格路线", config["route_name"])
    print_wrapped_field("生成文件", output_path_from_config(config, None).name)

    while True:
        try:
            answer = input("确认生成？请输入 Y/N：").strip().lower()
        except EOFError:
            fail("已退出，未生成行程单")
        if answer in ("y", "yes"):
            return True
        if answer in ("n", "no"):
            return False
        fail("已退出，未生成行程单")


def print_wrapped_field(label: str, value: Any) -> None:
    width = min(max(shutil.get_terminal_size((100, 24)).columns - 2, 72), 110)
    prefix = f"{label}："
    text = str(value or "")
    lines = wrap_display_width_with_prefix(prefix, text, width)
    for line in lines:
        print(line)


def wrap_display_width_with_prefix(prefix: str, text: str, width: int) -> list[str]:
    prefix_width = display_text_width(prefix)
    first_width = max(width - prefix_width, 20)
    subsequent_indent = " " * prefix_width

    first, rest = take_display_width(text, first_width)
    lines = [f"{prefix}{first}"]
    while rest:
        chunk, rest = take_display_width(rest, max(width - prefix_width, 20))
        lines.append(f"{subsequent_indent}{chunk}")
    return lines


def wrap_display_width(text: str, width: int) -> list[str]:
    lines: list[str] = []
    current = ""
    current_width = 0
    for char in text:
        char_width = display_width(char)
        if current and current_width + char_width > width:
            lines.append(current)
            current = char
            current_width = char_width
        else:
            current += char
            current_width += char_width
    lines.append(current)
    return lines or [""]


def display_width(char: str) -> int:
    return 2 if unicodedata.east_asian_width(char) in ("F", "W") else 1


def display_text_width(text: str) -> int:
    return sum(display_width(char) for char in text)


def take_display_width(text: str, width: int) -> tuple[str, str]:
    current = ""
    current_width = 0
    for index, char in enumerate(text):
        char_width = display_width(char)
        if current and current_width + char_width > width:
            return current, text[index:]
        current += char
        current_width += char_width
    return current, ""


def read_order_data(excel_path: Path) -> dict[str, Any]:
    if not excel_path.exists():
        fail(f"找不到订单 Excel：{excel_path}")

    workbook = load_workbook(excel_path, data_only=True)
    if SHEET_NAME not in workbook.sheetnames:
        fail(f"Excel 中找不到工作表：{SHEET_NAME}")

    sheet = workbook[SHEET_NAME]
    headers = [normalize_cell(cell.value) for cell in sheet[1]]
    missing = [column for column in REQUIRED_COLUMNS if column not in headers]
    if missing:
        fail("Excel 缺少必要列：" + "、".join(missing))

    indexes = {header: idx for idx, header in enumerate(headers)}
    passengers: list[dict[str, str]] = []
    route_name = ""
    travel_period = ""
    for row in sheet.iter_rows(min_row=2, values_only=True):
        if not any(value not in (None, "") for value in row):
            continue

        if "路线" in indexes and not route_name:
            route_name = normalize_cell(row[indexes["路线"]])
        if "出行时间" in indexes and not travel_period:
            travel_period = normalize_cell(row[indexes["出行时间"]])

        passenger = {column: normalize_cell(row[indexes[column]]) for column in REQUIRED_COLUMNS}
        if not passenger["旅客姓名"]:
            continue
        passengers.append(passenger)

    if not passengers:
        fail("Excel 没有可用旅客数据")

    start_date, end_date = parse_travel_period(travel_period)
    return {
        "passengers": passengers,
        "route_name": route_name,
        "start_date": start_date,
        "end_date": end_date,
    }


def parse_travel_period(value: str) -> tuple[str, str]:
    value = normalize_cell(value)
    if not value:
        return "", ""
    parts = re.split(r"\s*(?:~|至|到|-{2,}|—|－)\s*", value, maxsplit=1)
    if len(parts) != 2:
        return "", ""
    return normalize_date_text(parts[0]), normalize_date_text(parts[1])


def normalize_cell(value: Any) -> str:
    if value is None:
        return ""
    if isinstance(value, datetime):
        return value.strftime("%Y-%m-%d")
    if isinstance(value, date):
        return value.strftime("%Y-%m-%d")
    return str(value).strip()


def normalize_date_text(value: str) -> str:
    value = str(value).strip()
    for fmt in ("%Y-%m-%d", "%Y/%m/%d", "%Y.%m.%d"):
        try:
            return datetime.strptime(value, fmt).strftime("%Y-%m-%d")
        except ValueError:
            pass
    return value


def format_chinese_date(value: str) -> str:
    value = str(value).strip()
    try:
        parsed = datetime.strptime(value, "%Y-%m-%d")
    except ValueError:
        return value
    return f"{parsed.year} 年 {parsed.month:02d} 月 {parsed.day:02d} 日"


def sanitize_filename(value: str) -> str:
    cleaned = re.sub(r'[\\/:*?"<>|]+', "_", value).strip()
    return cleaned or "行程单.docx"


def set_cell_text(cell: Any, text: Any) -> None:
    text = "" if text is None else str(text)
    paragraphs = cell.paragraphs
    if not paragraphs:
        cell.text = text
        return

    first_paragraph = paragraphs[0]
    first_run = first_paragraph.runs[0] if first_paragraph.runs else first_paragraph.add_run()

    for paragraph in paragraphs[1:]:
        paragraph._element.getparent().remove(paragraph._element)

    for run in first_paragraph.runs:
        run.text = ""

    lines = text.splitlines() or [""]
    first_run.text = lines[0]

    for line in lines[1:]:
        first_run.add_break()
        first_run.add_text(line)


def fill_document(template_path: Path, output_path: Path, config: dict[str, Any], passengers: list[dict[str, str]]) -> None:
    if not template_path.exists():
        fail(f"找不到 DOCX 模板：{template_path}")

    document = Document(template_path)
    if not document.tables:
        fail("DOCX 模板中没有表格，无法填充")

    table = document.tables[0]
    if len(table.rows) < 32:
        fail("DOCX 模板表格行数不足，无法匹配当前行程单版式")

    set_cell_text(row_cell(table, 1, 1), config["route"])
    set_cell_text(row_cell(table, 2, 1), config["agency"])
    set_cell_text(row_cell(table, 2, 5), config["source"])
    set_cell_text(row_cell(table, 2, 7), len(passengers))
    set_cell_text(row_cell(table, 3, 1), format_chinese_date(config["start_date"]))
    set_cell_text(row_cell(table, 3, 5), config["leaders"])
    set_cell_text(row_cell(table, 4, 1), format_chinese_date(config["end_date"]))
    set_cell_text(row_cell(table, 4, 5), config["leader_phone"])
    center_cell_text(row_cell(table, 3, 5))
    center_cell_text(row_cell(table, 4, 5))
    set_cell_text(row_cell(table, 31, 1), config["operator"])
    set_cell_text(row_cell(table, 31, 3), config["operator_phone"])

    fill_passenger_table(row_cell(table, 5, 1), passengers)

    output_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(output_path)


def row_cell(table: Any, row_index: int, cell_index: int) -> Any:
    try:
        return table.rows[row_index].cells[cell_index]
    except IndexError:
        fail(f"DOCX 模板表格缺少单元格：第 {row_index + 1} 行，第 {cell_index + 1} 列")


def center_cell_text(cell: Any) -> None:
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    for paragraph in cell.paragraphs:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER


def fill_passenger_table(remarks_cell: Any, passengers: list[dict[str, str]]) -> None:
    if not remarks_cell.tables:
        fail("DOCX 模板备注区域缺少旅客明细表格")

    remarks_cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
    passenger_table = remarks_cell.tables[0]
    while len(passenger_table.rows) < len(passengers):
        passenger_table.add_row()

    for row_index, row in enumerate(passenger_table.rows):
        passenger = passengers[row_index] if row_index < len(passengers) else None
        values = passenger_columns(passenger) if passenger else ("", "", "", "")
        for cell, value in zip(row.cells, values):
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_text(cell, value)


def passenger_columns(passenger: dict[str, str]) -> tuple[str, str, str, str]:
    return (
        passenger["旅客姓名"],
        passenger["旅客证件类型"],
        passenger["旅客证件号码"],
        passenger["旅客手机号码"],
    )


def format_passenger(passenger: dict[str, str]) -> str:
    return " ".join(part for part in passenger_columns(passenger) if part)


def output_path_from_config(config: dict[str, Any], override: Path | None) -> Path:
    if override is not None:
        return resolve_path(override)
    default_name = default_output_name(config)
    name = sanitize_filename(str(config.get("output_name") or default_name))
    if not name.lower().endswith(".docx"):
        name += ".docx"
    return OUTPUT_DIR / name


def default_output_name(config: dict[str, Any]) -> str:
    start_date = str(config.get("start_date") or "").strip()
    route_name = str(config.get("route_name") or "").strip()
    date_code = filename_date_code(start_date)
    if date_code and route_name:
        return f"{date_code}{route_name}行程单.docx"
    if route_name:
        return f"{route_name}行程单_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
    return f"行程单_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"


def filename_date_code(start_date: str) -> str:
    try:
        parsed = datetime.strptime(start_date, "%Y-%m-%d")
    except ValueError:
        return ""
    return f"{parsed.year - 2020:02d}{parsed.month:02d}{parsed.day:02d}"


def main() -> None:
    args = parse_args()
    excel_path = resolve_path(args.excel) if args.excel else latest_order_file()
    template_path = resolve_path(args.template)
    config_path = resolve_path(args.config)

    order_data = read_order_data(excel_path)
    config = load_config(config_path)
    if order_data["start_date"]:
        config["start_date"] = order_data["start_date"]
    if order_data["end_date"]:
        config["end_date"] = order_data["end_date"]
    if order_data["route_name"]:
        config["route_name"] = order_data["route_name"]
    if not args.no_prompt:
        config = prompt_trip_fields(config)
    passengers = order_data["passengers"]
    output_path = output_path_from_config(config, args.output)

    fill_document(template_path, output_path, config, passengers)
    print(f"已生成：{output_path}")
    print(f"旅客人数：{len(passengers)}")


if __name__ == "__main__":
    try:
        main()
    except KeyboardInterrupt:
        sys.exit("已取消")
